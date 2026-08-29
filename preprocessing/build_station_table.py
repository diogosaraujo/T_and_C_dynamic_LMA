#!/usr/bin/env python3
"""The study's station table, as a Word document ready to paste.

One row per AmeriFlux station whose results reach the figures. That is the ERA5
fleet -- every station with an ERA5-Land pair in station_metrics.csv -- and a
station that does NOT also carry the full GCM set (historical + both SSPs, both
arms, all five models) is marked with an asterisk. Those are the 92 and 82 the
figures are built on; the counts come from the results themselves, never from
the 118-station site lists.

TWO TABLES, NOT ONE. Thirteen columns of 92 rows do not fit in 6.5 inches at
10 pt and stay readable -- the site-name and DOI columns alone need ~170 pt of
the 468 available. Table 1 is identity, location and citation; Table 2 is the
station-specific model parameters. Both are keyed by the same Site column, so
they read as one table split for width, which is standard for a supplement.
--layout single forces everything into one landscape-width table instead.

PARAMETERS COME FROM MOD_PARAM, NOT FROM THE FETCH TABLES. The MOD_PARAM file in
each station's run directory is what the model actually read, and it differs from
the fetched inputs in two ways that matter: ZR95_H is CAPPED at the soil column
depth (T&C aborts in Root_Fraction_General otherwise), and hc_H prefers AmeriFlux
BADM HEIGHTC over the Potapov/GEDI product where BADM has a value -- US-Ha2 is
21 m from BADM against 26 m from the raster. Reading the fetch outputs would put
numbers in the table that no run ever used.

WHICH PARAMETERS ARE STATION-SPECIFIC is taken from build_model_run.py's own
substitution list, not from a guess about what looks site-like:

    Zs / ms      soil column depth and layer count
    Psan/Pcla/Porg   texture, PER LAYER -- summarised here as a
                     THICKNESS-WEIGHTED profile mean, since the layers are graded
                     (10 mm at the top, 250 mm at the bottom) and a plain mean
                     would over-weight the thin surface layers
    Kbot         0.01 mm/h where bedrock is reported, NaN (free drainage) otherwise
    ZR95_H       rooting depth, capped at the column depth
    zatm         measurement height (BASE height, else max EC height, else hc+12)
    hc_H         canopy height
    Sl_H         SLA from the station's own LMA series. NOT TABULATED -- see
                 below. Carried in the CSV as lma_g_m2 = 1/(Sl_H*f_C).
    aSE_H        phenology switch = the vegetation-type column

LMA IS THE ONE STATION-SPECIFIC INPUT THAT IS NOT A PROPERTY OF THE STATION, so
Table 2 omits it by choice, not by oversight. MOD_PARAM's Sl_H is the FIXED arm's
temporal mean over its own forcing period; the dynamic arm varies around that
value every year, and both shift again under GCM forcing because the mean is then
taken over that model's period. Printing one number per station would invite it
to be read as "this station's LMA", and no such quantity exists. The value stays
in station_table.csv, where the arm and period it belongs to are explicit.

Everything else in MOD_PARAM is PFT-prescribed and identical across all stations
of a vegetation type, so it does not belong in a per-station table.

DOIs come from the cache written by fetch_ameriflux_dois.py. Sites publish a BASE
DOI always and a FLUXNET DOI only where ONEFlux has processed them; both are
shown where both exist. Elevation and the official site name come from one
request to the AmeriFlux sitemap, cached beside the DOIs.
"""
from __future__ import annotations

import argparse
import csv
import json
import re
import sys
import urllib.request
from pathlib import Path

import numpy as np
import pandas as pd

sys.path.insert(0, str(Path(__file__).resolve().parent))
from results_dir import NoResultsDir, resolve_out                 # noqa: E402

PREPROC = Path(__file__).resolve().parent
REPO_ROOT = PREPROC.parent
SITE_LISTS = [REPO_ROOT / "T&C" / "dynamic_lma_test" / "deciduous_ameriflux.csv",
              REPO_ROOT / "T&C" / "dynamic_lma_test" / "evergreen_ameriflux.csv"]
DOI_CACHE = PREPROC / "ameriflux_dois.csv"
META_CACHE = PREPROC / "ameriflux_site_meta.csv"
SITEMAP = "https://amfcdn.lbl.gov/api/v1/site_display/AmeriFlux"

# f_C, the carbon fraction of leaf dry mass. Sl = 1/(LMA*f_C), so LMA = 1/(Sl*f_C).
F_C = 0.5


class Missing(Exception):
    """A required input is absent. Never substituted with a plausible value."""


# ------------------------------------------------------------------- MOD_PARAM
def _num(txt: str, pattern: str):
    m = re.search(pattern, txt, re.M)
    if not m:
        return None, ""
    raw = m.group(1).strip()
    note = (m.group(2) or "").strip() if m.lastindex and m.lastindex >= 2 else ""
    if raw.lower() == "nan":
        return float("nan"), note
    try:
        return float(raw), note
    except ValueError:
        return None, note


def _vec(txt: str, name: str):
    m = re.search(rf"^{name}\s*=\s*\[([^\]]*)\]", txt, re.M)
    if not m:
        return None
    vals = [float(v) for v in re.split(r"[,\s]+", m.group(1).strip()) if v]
    return np.array(vals, float) if vals else None


def read_mod_param(path: Path) -> dict:
    """Every station-specific value in one MOD_PARAM, as the run used it."""
    t = path.read_text(errors="replace")
    out = {}

    zs = _vec(t, "Zs")
    if zs is not None and zs.size > 1:
        out["soil_depth_mm"] = float(zs[-1])
        out["ms"] = int(zs.size - 1)
        thick = np.diff(zs)
    else:
        thick = None

    # THICKNESS-WEIGHTED, not a plain mean: the mesh is graded, so the top 10 mm
    # layer and a 250 mm layer would otherwise count equally.
    for key, name in (("sand", "Psan_Zs"), ("clay", "Pcla_Zs"), ("org", "Porg_Zs")):
        v = _vec(t, name)
        if v is None:
            continue
        if thick is not None and thick.size == v.size:
            out[f"{key}_pct"] = float(100.0 * np.average(v, weights=thick))
        else:
            out[f"{key}_pct"] = float(100.0 * np.mean(v))

    kbot, _ = _num(t, r"^Kbot\s*=\s*([^;]+);(?:\s*%*\s*(.*))?$")
    if kbot is not None:
        # NaN is free drainage; a finite value is the bedrock conductivity.
        out["drainage"] = "free" if np.isnan(kbot) else "bedrock"
        out["Kbot"] = kbot
    for key, pat in (("zatm_m", r"^zatm\s*=\s*([0-9.eE+-]+)\s*;(?:\s*%*\s*(.*))?$"),
                     ("zr95_mm", r"^ZR95_H\s*=\s*\[([0-9.eE+-]+)\]\s*;(?:\s*%*\s*(.*))?$"),
                     ("sl_h", r"^Sl_H\s*=\s*\[([0-9.eE+-]+)\]\s*;(?:\s*%*\s*(.*))?$"),
                     ("hc_m", r"^hc_H\s*(?:\([^)]*\))?\s*=\s*\[([0-9.eE+-]+)\]\s*;(?:\s*%*\s*(.*))?$"),
                     ("ase", r"^aSE_H\s*=\s*\[([0-9]+)\]\s*;(?:\s*%*\s*(.*))?$")):
        v, note = _num(t, pat)
        if v is not None:
            out[key] = v
            if note:
                out[f"{key}_src"] = re.sub(r"^\[[^\]]*\]\s*", "", note).strip()
    if out.get("sl_h"):
        # The study's input variable, recovered from the SLA the run was given.
        out["lma_g_m2"] = 1.0 / (out["sl_h"] * F_C)
    return out


def collect_params(model_run: Path | None, stations) -> pd.DataFrame:
    """One row per station from its MOD_PARAM, or an empty frame if none found."""
    if not model_run or not Path(model_run).is_dir():
        print("  note: MODEL_RUN is not set or not a directory, so no "
              "station-specific parameters can be read", file=sys.stderr)
        return pd.DataFrame(columns=["station"])
    want = set(stations)

    def rank(p: Path) -> tuple:
        """Prefer the ERA5-Land fixed arm, in that order of importance.

        THE ARM MATTERS FOR ONE VALUE. hc, zatm, ZR95, Zs and the texture are
        identical in every arm of a station, but Sl_H is not: the fixed arm's is
        1/(mean(LMA)*f_C) over that arm's own LMA series, so a GCM arm reports a
        GCM-period mean. Taking whichever file sorted first gave US-Wrc an LMA
        labelled "GFDL-ESM4 1985-2014 mean" while its neighbours got ERA5 means
        -- one column silently mixing two definitions. This table is built on the
        ERA5 fleet, so it reads the ERA5-Land fixed arm wherever one exists.
        """
        s = str(p).lower().replace("\\", "/")
        return (0 if "era5" in s else 1, 0 if "fixed" in s else 1, s)

    best: dict = {}
    for f in Path(model_run).glob("*/**/MOD_PARAM_*.m"):
        st = f.parent
        while st.parent != Path(model_run) and st.parent != st:
            st = st.parent
        sid = st.name
        if sid not in want:
            continue
        if sid not in best or rank(f) < rank(best[sid]):
            best[sid] = f

    rows = []
    for sid, f in sorted(best.items()):
        rec = read_mod_param(f)
        if rec:
            rec["station"] = sid
            rows.append(rec)
    print(f"  MOD_PARAM: {len(rows)}/{len(want)} stations read")
    non_era5 = [s for s, f in best.items() if "era5" not in str(f).lower()]
    if non_era5:
        print(f"  note: {len(non_era5)} station(s) have no ERA5-Land MOD_PARAM, "
              f"so their LMA is a GCM-period mean: {', '.join(sorted(non_era5)[:6])}",
              file=sys.stderr)
    return pd.DataFrame(rows) if rows else pd.DataFrame(columns=["station"])


def load_dump(results: Path):
    """(station_inputs, mod_param_values) from dump_station_inputs.py, or Nones.

    PREFERRED OVER PARSING MOD_PARAM HERE. The dump is produced on the cluster,
    where the run tree, the forcing .mat files, the uncapped root-depth table and
    the FLUXNET archives all exist; this table then formats those numbers
    anywhere. It also carries the two things no local parse can know -- whether a
    station's ZR95 was capped, and whether its tower record entered the flux
    comparison.
    """
    a = Path(results) / "station_inputs.csv"
    b = Path(results) / "mod_param_values.csv"
    A = pd.read_csv(a, low_memory=False) if a.is_file() else None
    B = pd.read_csv(b, low_memory=False) if b.is_file() else None
    if A is None:
        print(f"  note: {a.name} not found -- falling back to reading MOD_PARAM "
              f"directly, which cannot mark capped ZR95 or the tower subset",
              file=sys.stderr)
    else:
        extra = (f", {B['parameter'].nunique()} parameters from {b.name}"
                 if B is not None else "")
        print(f"  dump: {len(A)} stations from {a.name}{extra}")
    return A, B


# ------------------------------------------------------------------- PFT table
# The curated vegetation-parameter list carries the acronym, the full name and
# the units. Values are NOT taken from it -- they come from the MOD_PARAM files
# the runs used, and this supplies only the wording.
PFT_DOC = PREPROC / "tc_evergreen_pft_parameters.csv"


def pft_rows(long: pd.DataFrame, sites: pd.DataFrame, layers=("H", "both"),
             differing_only: bool = False):
    """(rows, notes) for the per-vegetation-type parameter table.

    A parameter qualifies only if it is CONSTANT ACROSS EVERY STATION of a
    vegetation type. That is the definition of "assigned per vegetation type",
    and testing it against the actual files is the point: it proves the claim
    rather than restating a list someone wrote down. Anything that varies
    station to station is site-specific by construction and belongs in Table 2.
    """
    if long is None or long.empty:
        raise Missing("mod_param_values.csv not found -- run "
                      "dump_station_inputs.py on the cluster first")
    if not PFT_DOC.is_file():
        raise Missing(f"{PFT_DOC.name} not found; it supplies the parameter "
                      f"descriptions and units")
    doc = pd.read_csv(PFT_DOC)
    doc = doc[doc["layer"].isin(layers)] if "layer" in doc.columns else doc
    meta = {r["variable"]: (str(r.get("description") or ""),
                            str(r.get("units") or ""))
            for _, r in doc.iterrows()}

    d = long.merge(sites[["station", "pft"]], on="station", how="left")
    d = d[d["parameter"].isin(meta) & d["pft"].isin(["deciduous", "evergreen"])]
    if d.empty:
        raise Missing("no curated vegetation parameter matched the MOD_PARAM dump")

    rows, varying = [], []
    for name in sorted(d["parameter"].unique()):
        g = d[d["parameter"] == name]
        vals, ok = {}, True
        for pft in ("deciduous", "evergreen"):
            u = sorted(set(g[g["pft"] == pft]["value"].astype(str)))
            if not u:
                vals[pft] = "—"
            elif len(u) == 1:
                vals[pft] = u[0]
            else:
                ok = False
                varying.append(f"{name} ({pft}: {len(u)} distinct)")
                break
        if not ok:
            continue
        if differing_only and vals["deciduous"] == vals["evergreen"]:
            continue
        desc, units = meta[name]
        rows.append([name, desc, units,
                     _tidy(vals["deciduous"]), _tidy(vals["evergreen"])])
    return rows, varying


def _list(s):
    """A MATLAB bracketed list as floats, or None.

    Distinct from _vec above, which searches a whole MOD_PARAM text for a
    named assignment. This one parses a value already extracted into the
    dump. Naming both _vec silently shadowed the first and broke the
    no-dump fallback path with a TypeError.
    """
    parts = [p for p in re.split(r"[,\s]+", str(s).strip().strip("[]")) if p]
    try:
        return np.array([float(p) for p in parts], float)
    except ValueError:
        return None


def layer_rows(long: pd.DataFrame, stations) -> tuple:
    """(blocks, max_layers) for the depth-resolved soil table.

    Each block is (station, [(label, [cell, ...]), ...]) -- four sub-rows per
    site: sand, clay, organic and the layer's BOTTOM depth.

    DEPTH IS IN CENTIMETRES, not the millimetres Zs stores. That is purely a
    width decision and worth being explicit about: mm reaches 5000, four digits,
    ~20.4 pt at 10 pt, which does not fit an 18-column table inside 6.5 inches.
    In cm the maximum is 500 and every cell is at most three characters, so the
    table stays at 10 pt. Sand and clay are whole percent; organic keeps one
    decimal because it spans 0-8% and rounding would collapse most of it to 0.

    Sites with fewer than max_layers layers get an em dash in the unused columns
    rather than a blank, so a short profile is visibly short rather than looking
    like missing data.
    """
    if long is None or long.empty:
        raise Missing("mod_param_values.csv not found -- run "
                      "dump_station_inputs.py on the cluster first")
    per: dict = {}
    for r in long.itertuples():
        per.setdefault(r.station, {})[r.parameter] = r.value

    blocks, widths = [], []
    for st in stations:
        p = per.get(st)
        if not p:
            continue
        zs = _list(p.get("Zs"))
        if zs is None or zs.size < 2:
            continue
        series = {}
        for label, name, scale, nd in (("Sand (%)", "Psan_Zs", 100.0, 0),
                                       ("Clay (%)", "Pcla_Zs", 100.0, 0),
                                       ("Org. (%)", "Porg_Zs", 100.0, 1)):
            v = _list(p.get(name))
            series[label] = (None if v is None else
                             [f"{x * scale:.{nd}f}" for x in v])
        # Layer k spans Zs[k-1]..Zs[k]; the bottom interface is what a reader
        # needs to place the layer, so Zs[1:] is the row, not Zs[0].
        series["Depth (cm)"] = [f"{x / 10.0:.0f}" for x in zs[1:]]
        n = max((len(v) for v in series.values() if v), default=0)
        widths.append(n)
        blocks.append((st, series))

    if not blocks:
        raise Missing("no station had a readable Zs / texture profile")
    maxl = max(widths)
    rows = []
    for st, series in blocks:
        rows.append((st, [(lab, (series.get(lab) or []) +
                                ["—"] * (maxl - len(series.get(lab) or [])))
                          for lab in ("Sand (%)", "Clay (%)", "Org. (%)",
                                      "Depth (cm)")]))
    return rows, maxl


def _tidy(v: str) -> str:
    """Trim MATLAB brackets and pointless trailing zeros for display."""
    s = str(v).strip()
    s = re.sub(r"^\[\s*|\s*\]$", "", s).strip()
    try:
        f = float(s)
    except ValueError:
        # MOD_PARAM writes some rates as unevaluated arithmetic: the evergreen
        # file has Klf_H = [1/30] where the deciduous one has [0.025]. Left as
        # text the two columns are not comparable, so simple expressions are
        # evaluated. The character whitelist is what keeps this from being an
        # arbitrary eval of file contents.
        if re.fullmatch(r"[0-9.\s+\-*/()eE]+", s):
            try:
                f = float(eval(s, {"__builtins__": {}}, {}))   # noqa: S307
            except (ArithmeticError, SyntaxError, TypeError, ValueError):
                return s
        else:
            return s if s else "—"
    # NaN is a real MATLAB value here (Kbot = NaN means free drainage), so it is
    # displayed, not turned into an em dash and not passed to int().
    if not np.isfinite(f):
        return "NaN" if np.isnan(f) else ("Inf" if f > 0 else "−Inf")
    if f == int(f) and abs(f) < 1e15:
        return str(int(f))
    # 4 significant figures: 1/30 as "0.0333" rather than "0.0333333", which is
    # false precision for a prescribed rate.
    return f"{f:.4g}"


# ------------------------------------------------------------------- metadata
def read_sites() -> pd.DataFrame:
    rows = []
    for p in SITE_LISTS:
        if not p.exists():
            raise Missing(f"site list not found: {p}")
        with p.open(newline="", encoding="utf-8-sig") as fh:
            for r in csv.DictReader(fh):
                sid = (r.get("StationID") or "").strip()
                if not sid:
                    continue
                try:
                    lat, lon = float(r["Lat"]), float(r["Lon"])
                except (KeyError, TypeError, ValueError):
                    lat = lon = float("nan")
                rows.append({"station": sid,
                             "name": (r.get("StationName") or "").strip(),
                             "lat": lat, "lon": lon,
                             "pft": (r.get("ForestType") or "").strip().lower(),
                             "igbp": (r.get("IGBP") or "").strip()})
    return pd.DataFrame(rows).drop_duplicates("station")


def site_meta(refresh: bool = False) -> pd.DataFrame:
    """Official site name and elevation, from ONE sitemap request, cached."""
    if META_CACHE.is_file() and not refresh:
        return pd.read_csv(META_CACHE)
    print("  fetching the AmeriFlux sitemap (one request, all sites)")
    rq = urllib.request.Request(SITEMAP, headers={"User-Agent": "Mozilla/5.0"})
    data = json.load(urllib.request.urlopen(rq, timeout=120))
    rows = []
    for s in data:
        loc = s.get("GRP_LOCATION") or {}
        rows.append({"station": s.get("SITE_ID", ""),
                     "af_name": s.get("SITE_NAME", ""),
                     "af_igbp": s.get("IGBP", ""),
                     "elev_m": loc.get("LOCATION_ELEV", "")})
    d = pd.DataFrame(rows)
    d.to_csv(META_CACHE, index=False)
    print(f"  -> {META_CACHE}  ({len(d)} sites)")
    return d


def fleets(results: Path):
    """(era5 stations, stations complete in all four datasets)."""
    p = Path(results) / "station_metrics.csv"
    if not p.is_file():
        raise Missing(f"{p} not found -- run station_metrics.py first")
    m = pd.read_csv(p, low_memory=False)
    m = m[(m["freq"] == "annual") & (m["subset"] == "all")]
    per = {ds: set(g["station"].unique()) for ds, g in m.groupby("dataset")}
    if "era5" not in per:
        raise Missing("station_metrics.csv has no era5 rows")
    return sorted(per["era5"]), set.intersection(*per.values())


# ---------------------------------------------------------------------- build
def assemble(results: Path, model_run: Path | None, refresh: bool) -> pd.DataFrame:
    era5, gcm = fleets(results)
    print(f"  fleets: {len(era5)} ERA5, {len(gcm)} with the full GCM set, "
          f"{len(set(era5) - gcm)} asterisked")

    d = pd.DataFrame({"station": era5})
    d["gcm"] = d["station"].isin(gcm)
    d = d.merge(read_sites(), on="station", how="left")
    d = d.merge(site_meta(refresh), on="station", how="left")

    if DOI_CACHE.is_file():
        doi = pd.read_csv(DOI_CACHE).fillna("")
        d = d.merge(doi, on="station", how="left")
    else:
        print(f"  note: {DOI_CACHE.name} not found -- run "
              f"fetch_ameriflux_dois.py; the DOI columns will be empty",
              file=sys.stderr)
        for c in ("doi_base", "doi_fluxnet", "citation_base", "citation_fluxnet"):
            d[c] = ""
    d[["doi_base", "doi_fluxnet"]] = d[["doi_base", "doi_fluxnet"]].fillna("")

    dump, long = load_dump(results)
    if dump is not None:
        keep = [c for c in dump.columns if c not in ("in_gcm",)]
        d = d.merge(dump[keep], on="station", how="left")
        # Zbas as the model received it beats the registry value: they can
        # differ, and only one of them was in the run.
        if "mat_Zbas" in d.columns:
            n = int(d["mat_Zbas"].notna().sum())
            reg = pd.to_numeric(d["elev_m"], errors="coerce")
            gap = (reg - d["mat_Zbas"]).abs()
            far = d[gap > 5]
            d["elev_m"] = d["mat_Zbas"].where(d["mat_Zbas"].notna(), reg)
            d["elev_src"] = np.where(d["mat_Zbas"].notna(),
                                     "forcing .mat", "AmeriFlux registry")
            print(f"  elevation: {n}/{len(d)} from the forcing .mat"
                  f"{f', {len(far)} differ from the registry by >5 m' if len(far) else ''}")
            for _, r in far.head(5).iterrows():
                print(f"    {r['station']}: .mat {r['mat_Zbas']:.0f} m vs "
                      f"registry {reg[r.name]:.0f} m")
    else:
        long = None
        par = collect_params(model_run, era5)
        d = d.merge(par, on="station", how="left")
        d["elev_src"] = "AmeriFlux registry"
    for c in ("zr95_capped", "tower_used"):
        if c not in d.columns:
            d[c] = np.nan
    d.attrs["long"] = long

    # The site lists and AmeriFlux must agree on vegetation type. aSE_H is the
    # switch the model ran with, so a disagreement is a real inconsistency.
    if "ase" in d.columns:
        exp = d["pft"].map({"evergreen": 0.0, "deciduous": 1.0})
        bad = d[d["ase"].notna() & exp.notna() & (d["ase"] != exp)]
        if len(bad):
            print(f"  WARNING: {len(bad)} station(s) whose aSE_H disagrees with "
                  f"the site list: {', '.join(bad['station'])}", file=sys.stderr)

    miss = d[d["lat"].isna() | d["lon"].isna()]
    if len(miss):
        print(f"  WARNING: {len(miss)} station(s) with no coordinates: "
              f"{', '.join(miss['station'])}", file=sys.stderr)
    return d.sort_values("station").reset_index(drop=True)


def fmt(v, nd=0, dash="—"):
    if v is None or (isinstance(v, float) and not np.isfinite(v)):
        return dash
    try:
        return f"{float(v):.{nd}f}"
    except (TypeError, ValueError):
        return str(v) if str(v).strip() else dash


# TABLE 1 -- identity, location, citation.
T1 = [("Site", 0.62), ("Site name", 1.62), ("Lat (°N)", 0.55), ("Lon (°E)", 0.62),
      ("Veg.", 0.48), ("Elev. (m)", 0.55), ("AmeriFlux DOI (10.17190/…)", 2.06)]
# TABLE 2 -- the station-specific model parameters.
# LMA IS DELIBERATELY NOT A COLUMN. It is the one station-specific input whose
# value is not a property of the station: MOD_PARAM's Sl_H is the FIXED arm's
# temporal mean, the dynamic arm varies around it yearly, and both change again
# under GCM forcing because the mean is taken over that model's own period. A
# single number in a station table would therefore be read as "this station's
# LMA" when no such quantity exists. It stays in station_table.csv, where the
# arm and period are explicit.
# Texture is NOT here. A profile mean duplicated what Table 3 shows properly,
# and the two disagreed in appearance even though they agreed arithmetically:
# US-Bar reads 81% sand as a thickness-weighted mean while its surface layers
# are 59%. One depth-resolved statement is better than a summary that invites
# the surface to be read off it.
T2 = [("Site", 0.62), ("Veg.", 0.48), ("hc (m)", 0.6), ("zatm (m)", 0.68),
      ("ZR95 (mm)", 0.78), ("Soil depth (mm)", 0.95), ("Soil Layers", 0.7)]
# TABLE 3 -- parameters prescribed per vegetation type.
T3 = [("Parameter", 0.85), ("Description", 2.7), ("Units", 0.9),
      ("Deciduous", 1.0), ("Evergreen", 1.05)]


def _flag(v) -> bool:
    """True only for a genuine true. NaN IS NOT TRUE.

    bool(float("nan")) is True in Python, and missing columns are filled with
    NaN, so a plain bool() marked all 92 stations as capped -- a footnote saying
    every station's rooting depth was cut, which is both wrong and exactly the
    kind of wrong that looks deliberate. v == v is False for NaN.
    """
    if v is None:
        return False
    if isinstance(v, str):
        return v.strip().lower() in ("true", "1", "yes")
    try:
        return bool(v) and v == v
    except (TypeError, ValueError):
        return False


def markers(r) -> tuple:
    """(site id, superscript) -- the fleet asterisk plus the two data flags.

    * not carried through the GCM experiments
    a ZR95 capped at the soil column depth
    b tower record entered the GPP/ET/LE/H comparison
    """
    m = []
    if not r["gcm"]:
        m.append("*")
    if _flag(r.get("zr95_capped")):
        m.append("a")
    if _flag(r.get("tower_used")):
        m.append("b")
    return str(r["station"]), ",".join(m)


def rows_t1(d: pd.DataFrame) -> list:
    out = []
    for _, r in d.iterrows():
        dois = [x.replace("10.17190/", "") for x in
                (str(r.get("doi_base", "")), str(r.get("doi_fluxnet", "")))
                if x and x != "nan"]
        name = str(r.get("af_name") or r.get("name") or "")
        out.append([markers(r), name,
                    fmt(r["lat"], 4), fmt(r["lon"], 4),
                    {"deciduous": "Dec.", "evergreen": "Eve."}.get(r["pft"], "—"),
                    fmt(r.get("elev_m"), 0), "\n".join(dois) if dois else "—"])
    return out


def rows_t2(d: pd.DataFrame) -> list:
    out = []
    for _, r in d.iterrows():
        depth = fmt(r.get("soil_depth_mm"), 0)
        # A dagger marks a bedrock lower boundary (Kbot = 0.01 mm/h); everything
        # else drains freely. It is a per-station choice, so it belongs here.
        free = r.get("kbot_free")
        bedrock = (str(r.get("drainage", "")) == "bedrock" or free is False)
        if bedrock and depth != "—":
            depth += " †"
        out.append([markers(r),
                    {"deciduous": "Dec.", "evergreen": "Eve."}.get(r["pft"], "—"),
                    fmt(r.get("hc_m"), 1), fmt(r.get("zatm_m"), 1),
                    fmt(r.get("zr95_mm"), 0), depth, fmt(r.get("ms"), 0)])
    return out


def write_docx(d: pd.DataFrame, out: Path, layout: str, pt: float,
               t3=None, layers=None, layer_font: float = 10.0) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    doc = Document()
    sec = doc.sections[0]
    if layout == "single":
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
        usable = 9.0
    else:
        usable = 6.5
    sec.left_margin = sec.right_margin = Inches(1.0)

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(pt)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.space_before = Pt(0)

    n_ast = int((~d["gcm"]).sum())

    def add_table(title, caption, spec, rows):
        h = doc.add_paragraph()
        h.add_run(title).bold = True
        doc.add_paragraph(caption)
        scale = usable / sum(w for _, w in spec)
        t = doc.add_table(rows=1, cols=len(spec))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        for j, (head, w) in enumerate(spec):
            c = t.rows[0].cells[j]
            c.text = ""
            run = c.paragraphs[0].add_run(head)
            run.bold = True
            run.font.size = Pt(pt)
            c.width = Inches(w * scale)
        for row in rows:
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].width = Inches(spec[j][1] * scale)
                para = cells[j].paragraphs[0]
                # A (base, superscript) pair renders as two runs so the markers
                # sit above the baseline; anything else is plain text.
                base, sup = val if isinstance(val, tuple) else (val, "")
                # Newlines inside a cell must become real line breaks; assigning
                # "a\nb" to cell.text renders as one run with a stray character.
                for k, line in enumerate(str(base).split("\n")):
                    if k:
                        para.add_run().add_break()
                    r = para.add_run(line)
                    r.font.size = Pt(pt)
                if sup:
                    r = para.add_run(sup)
                    r.font.size = Pt(pt)
                    r.font.superscript = True
                if j >= 2:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    def set_cell_margins(table, pt_margin=1.5):
        """Shrink the table's cell padding.

        Word defaults to 0.08 in per side, which is 11.5 pt of the ~21 pt each
        column gets in an 18-column table -- enough room for two characters.
        python-docx has no API for this, so it is set on w:tblCellMar directly.
        """
        tw = str(int(pt_margin * 20))                      # points -> twips
        mar = OxmlElement("w:tblCellMar")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), tw)
            el.set(qn("w:type"), "dxa")
            mar.append(el)
        table._tbl.tblPr.append(mar)

    def add_layer_table(title, caption, blocks, maxl, fpt):
        h = doc.add_paragraph()
        h.add_run(title).bold = True
        doc.add_paragraph(caption)
        ncol = 2 + maxl
        w_site, w_lab = 0.55, 0.62
        w_layer = (usable - w_site - w_lab) / maxl
        t = doc.add_table(rows=1, cols=ncol)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        set_cell_margins(t)
        heads = ["Site", ""] + [str(i) for i in range(1, maxl + 1)]
        for j, head in enumerate(heads):
            c = t.rows[0].cells[j]
            c.text = ""
            r = c.paragraphs[0].add_run(head)
            r.bold = True
            r.font.size = Pt(fpt)
            c.width = Inches(w_site if j == 0 else w_lab if j == 1 else w_layer)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for st, subrows in blocks:
            first = None
            for k, (label, vals) in enumerate(subrows):
                cells = t.add_row().cells
                for j in range(ncol):
                    cells[j].width = Inches(w_site if j == 0
                                            else w_lab if j == 1 else w_layer)
                if k == 0:
                    first = cells[0]
                    run = cells[0].paragraphs[0].add_run(st)
                    run.font.size = Pt(fpt)
                else:
                    # Vertically merge the Site cell down the block so the id
                    # appears once per station rather than four times.
                    first = first.merge(cells[0])
                r = cells[1].paragraphs[0].add_run(label)
                r.font.size = Pt(fpt)
                for j, v in enumerate(vals):
                    p = cells[2 + j].paragraphs[0]
                    rr = p.add_run(str(v))
                    rr.font.size = Pt(fpt)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    doc.add_paragraph()
    if layout == "single":
        add_table("Table 1. AmeriFlux stations used in this study.",
                  f"Sites with an ERA5-Land simulation pair (n = {len(d)}). "
                  f"Asterisk: not carried through the GCM experiments "
                  f"(n = {n_ast}).",
                  T1[:-1] + T2[2:] + T1[-1:],
                  [a[:-1] + b[2:] + a[-1:]
                   for a, b in zip(rows_t1(d), rows_t2(d))])
    else:
        n_cap = int(d["zr95_capped"].map(_flag).sum())
        n_tow = int(d["tower_used"].map(_flag).sum())
        # d.get() on a MISSING column returns None, and None == "..." is a plain
        # bool with no .any(), which crashed job 41104 after both steps had done
        # their work. Check the column exists before comparing.
        elev_src = ("the model forcing files"
                    if ("elev_src" in d.columns
                        and (d["elev_src"] == "forcing .mat").any())
                    else "the AmeriFlux site registry")
        foot = (f"* the {n_ast} stations that do not also carry the complete "
                f"GCM set (historical and both SSP scenarios, five models, both "
                f"arms), and so appear only in the ERA5-Land results. "
                f"ᵃ rooting depth capped at the soil column depth (n = {n_cap}). "
                f"ᵇ tower measurements entered the GPP, ET, LE and H comparison "
                f"(n = {n_tow}).")
        add_table(
            "Table 1. AmeriFlux stations used in this study.",
            f"All {len(d)} stations with an ERA5-Land simulation pair. "
            f"Elevation is the value used by the model, read from {elev_src}. "
            f"DOIs share the prefix 10.17190/; where two are listed, the first "
            f"is the AmeriFlux BASE product and the second the AmeriFlux "
            f"FLUXNET (ONEFlux) product. {foot}",
            T1, rows_t1(d))
        add_table(
            "Table 2. Station-specific model parameters.",
            f"Values as read by the model from each station's MOD_PARAM file. "
            f"All other T&C parameters are prescribed per vegetation type "
            f"(Table 4) and are identical across stations of that type. ZR95 is "
            f"the Schenk & Jackson rooting depth, capped at the soil column "
            f"depth where it would otherwise exceed it (ᵃ, n = {n_cap}); T&C "
            f"aborts in Root_Fraction_General without that cap. Soil layers is "
            f"the number of layers in the station's mesh; the layer depths and "
            f"the texture of each are given in Table 3. † marks a bedrock "
            f"lower boundary (Kbot = 0.01 mm h⁻¹); all others drain freely. "
            f"{foot}",
            T2, rows_t2(d))
        if layers is not None:
            blocks, maxl = layers
            add_layer_table(
                "Table 3. Depth-resolved soil texture and layer geometry.",
                f"The soil profile each station was run with, layer by layer. "
                f"Saxton & Rawls is applied per layer, so these are the values "
                f"behind the hydraulic and thermal properties, and the "
                f"profile-mean texture in Table 2 is a summary of them. Depth is "
                f"the BOTTOM of each layer, in cm. Column count is the deepest "
                f"profile in the fleet ({maxl} layers); an em dash marks layers a "
                f"shallower site does not have. Silt is not tabulated because it "
                f"is not a model input — T&C derives it internally as "
                f"1 − sand − clay − organic.",
                blocks, maxl, layer_font)
        if t3 is not None:
            rows3, varying = t3
            add_table(
                "Table 4. Parameters prescribed per vegetation type.",
                "Every parameter that is constant across all stations of a "
                "vegetation type, verified station by station against the "
                "MOD_PARAM files rather than assumed from a list. Values are "
                "the High (canopy) vegetation layer; the Low layer is disabled "
                "at every site. Parameters that vary between stations are "
                "site-specific and appear in Table 2 instead.",
                T3, rows3)
            if varying:
                p = doc.add_paragraph()
                p.add_run(
                    f"Not tabulated: {len(varying)} curated parameter(s) were "
                    f"not constant within a vegetation type and are therefore "
                    f"not prescribed per type — " + "; ".join(varying[:8]) +
                    ("; …" if len(varying) > 8 else "")).italic = True

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--results", type=Path, default=None)
    ap.add_argument("--model-run", type=Path, default=None,
                    help="run tree holding <station>/**/MOD_PARAM_*.m; "
                         "defaults to $MODEL_RUN")
    ap.add_argument("--out", type=Path, default=None,
                    help="output .docx (default station_table.docx under $TC_RESULTS)")
    ap.add_argument("--csv", type=Path, default=None,
                    help="also write the full table, every column, as CSV")
    ap.add_argument("--layout", default="split", choices=["split", "single"],
                    help="split: two 6.5in portrait tables (default). "
                         "single: one landscape table")
    ap.add_argument("--font", type=float, default=10.0)
    ap.add_argument("--refresh-meta", action="store_true")
    ap.add_argument("--no-layer-table", action="store_true",
                    help="skip Table 3, the depth-resolved soil profile")
    ap.add_argument("--layer-font", type=float, default=10.0,
                    help="point size for Table 3, which is much denser than the "
                         "others; drop to 9 if depths are shown in mm")
    ap.add_argument("--no-pft-table", action="store_true",
                    help="skip Table 3, the per-vegetation-type parameters")
    ap.add_argument("--pft-differing-only", action="store_true",
                    help="Table 3 lists only parameters whose deciduous and "
                         "evergreen values differ")
    a = ap.parse_args(argv)

    import os
    mr = a.model_run or os.environ.get("MODEL_RUN") or os.environ.get("TC_MODEL_RUN")
    try:
        results = Path(a.results or resolve_out(".", create=False))
        out = Path(a.out) if a.out else resolve_out("station_table.docx")
    except NoResultsDir as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1

    try:
        d = assemble(results, Path(mr) if mr else None, a.refresh_meta)
    except Missing as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1

    layers = None
    if not a.no_layer_table and a.layout != "single":
        try:
            layers = layer_rows(d.attrs.get("long"), list(d["station"]))
            print(f"  Table 3: {len(layers[0])} stations, deepest profile "
                  f"{layers[1]} layers")
        except Missing as e:
            print(f"  note: Table 3 omitted -- {e}", file=sys.stderr)

    t3 = None
    if not a.no_pft_table and a.layout != "single":
        try:
            t3 = pft_rows(d.attrs.get("long"), d, differing_only=a.pft_differing_only)
            print(f"  Table 4: {len(t3[0])} parameters constant within a "
                  f"vegetation type"
                  + (f", {len(t3[1])} rejected as varying" if t3[1] else ""))
        except Missing as e:
            print(f"  note: Table 3 omitted -- {e}", file=sys.stderr)

    write_docx(d, out, a.layout, a.font, t3, layers, a.layer_font)
    print(f"\n-> {out}")
    if a.csv:
        p = Path(a.csv) if Path(a.csv).is_absolute() else resolve_out(a.csv)
        d.to_csv(p, index=False)
        print(f"-> {p}")

    have = int(d["hc_m"].notna().sum()) if "hc_m" in d.columns else 0
    print(f"   {len(d)} stations, {int((~d['gcm']).sum())} asterisked")
    print(f"   parameters from MOD_PARAM: {have}/{len(d)}")
    print(f"   BASE DOI: {int((d['doi_base'].astype(str).str.len() > 3).sum())}"
          f"/{len(d)}   FLUXNET DOI: "
          f"{int((d['doi_fluxnet'].astype(str).str.len() > 3).sum())}/{len(d)}")
    if have < len(d):
        print(f"   ! {len(d) - have} station(s) have no MOD_PARAM under "
              f"{mr or '$MODEL_RUN (unset)'}; their parameter cells are '—'.\n"
              f"     Run this on the cluster to fill them.", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
